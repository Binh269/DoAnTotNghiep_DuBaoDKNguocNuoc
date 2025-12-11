import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt

# Hàm đọc dữ liệu từ file Excel hoặc CSV
def load_data_from_file(file_path):
    try:
        if file_path.endswith('.xlsx'):
            df = pd.read_excel(file_path)
        elif file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            raise ValueError("Định dạng file không được hỗ trợ. Vui lòng sử dụng file .xlsx hoặc .csv")

        if len(df.columns) != 2:
            raise ValueError("File phải có đúng 2 cột: thời gian và giá trị")

        df.columns = ['Time', 'Value']

        is_year_only = True
        for t in df['Time']:
            try:
                int(t)
            except (ValueError, TypeError):
                is_year_only = False
                break

        if is_year_only:
            df['Time'] = df['Time'].astype(str)
        else:
            df['Time'] = pd.to_datetime(df['Time'], errors='coerce')
            if df['Time'].isna().any():
                raise ValueError(
                    "Không thể phân tích cột thời gian. Vui lòng đảm bảo định dạng là ngày tháng (YYYY-MM-DD)")

        if not np.issubdtype(df['Value'].dtype, np.number):
            raise ValueError("Cột giá trị phải chứa dữ liệu số")

        return df, is_year_only
    except Exception as e:
        print(f"Lỗi khi đọc file: {e}")
        return None, None

# Bước 1: Phân vùng tập hợp luận với D1, D2 tối ưu và số khoảng
def define_universe_and_intervals(data, D1, D2, num_intervals):
    D_min, D_max = min(data), max(data)
    U = [D_min - D1, D_max + D2]
    interval_width = (U[1] - U[0]) / num_intervals
    intervals = [[U[0] + i * interval_width, U[0] + (i + 1) * interval_width] for i in range(num_intervals)]
    midpoints = [(u[0] + u[1]) / 2 for u in intervals]
    return intervals, midpoints

# Hiển thị khoảng và phân bố dữ liệu
def display_intervals_and_distribution(data, intervals):
    print("\nCác khoảng phân vùng:")
    for i, (low, high) in enumerate(intervals):
        print(f"Khoảng {i+1}: [{low:.2f}, {high:.2f}]")

    distribution = [0] * len(intervals)
    for value in data:
        for i, (low, high) in enumerate(intervals):
            if low <= value <= high:
                distribution[i] += 1
                break

    plt.figure(figsize=(10, 6))
    plt.bar(range(1, len(intervals) + 1), distribution, tick_label=[f"K{i+1}" for i in range(len(intervals))])
    plt.title("Phân bố dữ liệu trong các khoảng")
    plt.xlabel("Khoảng")
    plt.ylabel("Số lượng giá trị")
    plt.grid(True)
    plt.show()

# Bước 2: Định nghĩa tập hợp mờ động
def define_fuzzy_sets(num_intervals):
    fuzzy_sets = {f'A_{i+1}': i for i in range(num_intervals)}
    return fuzzy_sets

# Vẽ đồ thị độ thuộc hàm Gauss
def plot_gaussian_membership(intervals, midpoints, sigmas):
    x_values = np.linspace(min([l for l, _ in intervals]) - 10, max([h for _, h in intervals]) + 10, 500)
    plt.figure(figsize=(12, 6))
    for i, (m, s) in enumerate(zip(midpoints, sigmas)):
        mu = [np.exp(-((x - m) ** 2) / (2 * (s ** 2))) for x in x_values]
        plt.plot(x_values, mu, label=f'A_{i+1} (m={m:.2f}, σ={s:.2f})')
    plt.title("Độ thuộc hàm Gauss của các tập mờ")
    plt.xlabel("Giá trị")
    plt.ylabel("Độ thuộc (μ)")
    plt.legend()
    plt.grid(True)
    plt.show()

# Bước 3: Mờ hóa dữ liệu với sigma tối ưu
def fuzzify_data_gaussian_weighted(data, intervals, fuzzy_sets, sigmas):
    fuzzified = []
    midpoints = [(l + u) / 2 for l, u in intervals]
    for x in data:
        mu = [np.exp(-((x - m) ** 2) / (2 * (s ** 2))) for m, s in zip(midpoints, sigmas)]
        total_mu = sum(mu)
        if total_mu > 0:
            fx = sum(mu_i * m for mu_i, m in zip(mu, midpoints)) / total_mu
        else:
            fx = midpoints[np.argmin([abs(x - m) for m in midpoints])]
        closest_idx = np.argmin([abs(fx - m) for m in midpoints])
        best_set = [key for key, value in fuzzy_sets.items() if value == closest_idx][0]
        fuzzified.append(best_set)
    return fuzzified

# Thêm mới: Hàm vẽ biểu đồ kết quả mờ hóa
def plot_fuzzified_data(times, values, fuzzified_data, is_year_only):
    if len(times) != len(values) or len(values) != len(fuzzified_data):
        print(f"Warning: Length mismatch - times: {len(times)}, values: {len(values)}, fuzzified_data: {len(fuzzified_data)}")
        min_len = min(len(times), len(values), len(fuzzified_data))
        times = times[:min_len]
        values = values[:min_len]
        fuzzified_data = fuzzified_data[:min_len]

    plt.figure(figsize=(12, 6))
    plt.plot(times, values, label='Giá trị thực', marker='o', color='blue')

    # Thêm nhãn tập mờ tại mỗi điểm
    for i, (t, v, f) in enumerate(zip(times, values, fuzzified_data)):
        plt.text(t, v + 0.02 * (max(values) - min(values)), f, fontsize=9, ha='center', color='red')

    plt.title("Kết quả Mờ hóa Chuỗi Thời gian")
    plt.xlabel("Thời gian")
    plt.ylabel("Giá trị")
    plt.legend()
    plt.grid(True)
    plt.xticks(rotation=45)

    # Định dạng trục x
    if is_year_only:
        plt.gca().set_xticks(range(0, len(times), max(1, len(times) // 10)))
    else:
        plt.gca().xaxis.set_major_formatter(plt.matplotlib.dates.DateFormatter('%Y-%m-%d'))

    plt.tight_layout()
    plt.show()

# Bước 4: Xác định FLRs
def identify_flrs(fuzzified_data, order):
    flrs = []
    for i in range(len(fuzzified_data) - order):
        current = tuple(fuzzified_data[i:i + order])
        next_state = fuzzified_data[i + order]
        flrs.append((current, next_state))
    return flrs

# Bước 5: Thiết lập Nhóm Quan hệ Mờ Phụ thuộc Thời gian theo từng năm
def establish_time_dependent_flrgs(fuzzified_data, order):
    yearly_flrgs = []
    for t in range(order, len(fuzzified_data)):
        flrgs = {}
        # Duyệt qua tất cả FLRs từ đầu đến thời điểm t
        for i in range(t - order + 1):
            current = tuple(fuzzified_data[i:i + order])
            next_state = fuzzified_data[i + order]
            if current not in flrgs:
                flrgs[current] = []
            flrgs[current].append(next_state)
        yearly_flrgs.append(flrgs)
    return yearly_flrgs

# Bước 6: Giải mờ và dự báo
def defuzzify_forecast_time_weighted(fuzzified_data, yearly_flrgs, midpoints, order):
    forecasts = [None] * order
    for t in range(order, len(fuzzified_data)):
        current = tuple(fuzzified_data[t - order:t])
        flrgs = yearly_flrgs[t - order]
        if current in flrgs:
            next_states = flrgs[current]
            weights = list(range(1, len(next_states) + 1))
            indices = [int(state.split('_')[1]) - 1 for state in next_states]
            mid_values = [midpoints[idx] for idx in indices]
            numerator = sum(w * m for w, m in zip(weights, mid_values))
            denominator = sum(weights)
            forecast = numerator / denominator if denominator > 0 else mid_values[-1]
            forecasts.append(forecast)
        else:
            idx = int(fuzzified_data[t - 1].split('_')[1]) - 1
            forecasts.append(midpoints[idx])
    return forecasts

# Hàm tính MSE để tối ưu D1, D2, và sigmas
def evaluate_forecast(individual, data, num_intervals, order, return_forecasts=False):
    D1, D2 = individual[0], individual[1]
    sigmas = individual[2:]
    if D1 <= 0 or D2 <= 0 or any(s <= 0 for s in sigmas):
        return float('inf') if not return_forecasts else (float('inf'), [])

    intervals, midpoints = define_universe_and_intervals(data, D1, D2, num_intervals)
    fuzzy_sets = define_fuzzy_sets(num_intervals)
    fuzzified_data = fuzzify_data_gaussian_weighted(data, intervals, fuzzy_sets, sigmas)
    flrs = identify_flrs(fuzzified_data, order)
    yearly_flrgs = establish_time_dependent_flrgs(fuzzified_data, order)
    forecasts = defuzzify_forecast_time_weighted(fuzzified_data, yearly_flrgs, midpoints, order)
    valid_indices = [i for i, f in enumerate(forecasts) if f is not None]
    actual = [data[i] for i in valid_indices]
    predicted = [forecasts[i] for i in valid_indices]
    mse = np.mean((np.array(actual) - np.array(predicted)) ** 2)
    if return_forecasts:
        return mse, forecasts
    return mse

# Tối ưu D1, D2, và sigmas bằng PSO
def optimize_D1_D2_sigmas(data, num_intervals, order, num_particles=50, max_iter=100):
    w_max, w_min = 0.9, 0.4
    c1, c2 = 2.0, 2.0
    v_max = 100
    D_bounds = (1, 1000)

    D_min, D_max = min(data), max(data)
    U_approx = [D_min - 500, D_max + 500]
    interval_width = (U_approx[1] - U_approx[0]) / num_intervals
    sigma_bounds = (0.1 * interval_width, 1.0 * interval_width)

    dim = 2 + num_intervals
    particles = np.zeros((num_particles, dim))
    particles[:, 0:2] = np.random.uniform(D_bounds[0], D_bounds[1], (num_particles, 2))
    particles[:, 2:] = np.random.uniform(sigma_bounds[0], sigma_bounds[1], (num_particles, num_intervals))

    velocities = np.random.uniform(-v_max, v_max, (num_particles, dim))
    personal_best = particles.copy()
    personal_best_scores = np.array([evaluate_forecast(p, data, num_intervals, order) for p in particles])
    global_best_idx = np.argmin(personal_best_scores)
    global_best = personal_best[global_best_idx].copy()
    global_best_score = personal_best_scores[global_best_idx]
    no_improve_count = 0
    max_no_improve = 20

    for iter in range(max_iter):
        w = w_max - (w_max - w_min) * iter / max_iter
        for i in range(num_particles):
            r1, r2 = np.random.rand(2)
            velocities[i] = (w * velocities[i] +
                             c1 * r1 * (personal_best[i] - particles[i]) +
                             c2 * r2 * (global_best - particles[i]))
            velocities[i] = np.clip(velocities[i], -v_max, v_max)
            particles[i] += velocities[i]
            particles[i, 0:2] = np.clip(particles[i, 0:2], D_bounds[0], D_bounds[1])
            particles[i, 2:] = np.clip(particles[i, 2:], sigma_bounds[0], sigma_bounds[1])
            score = evaluate_forecast(particles[i], data, num_intervals, order)
            if score < personal_best_scores[i]:
                personal_best[i] = particles[i].copy()
                personal_best_scores[i] = score
                if score < global_best_score:
                    global_best = particles[i].copy()
                    global_best_score = score
                    no_improve_count = 0
                else:
                    no_improve_count += 1
        if no_improve_count >= max_no_improve:
            print(f"Dừng sớm sau {iter+1} lần lặp vì không cải thiện.")
            break

    return global_best[0], global_best[1], global_best[2:], global_best_score

# Hàm xuất bảng ra file Word
def export_to_word(flrs, yearly_flrgs, results, times, is_year_only, output_file, order):
    doc = Document()

    doc.add_heading('Kết quả Dự báo Chuỗi Thời gian Mờ', 0)

    doc.add_heading('Bảng Quan hệ Mờ (FLRs)', level=1)
    flr_table = doc.add_table(rows=1 + len(flrs), cols=3)
    flr_table.style = 'Table Grid'
    hdr_cells = flr_table.rows[0].cells
    hdr_cells[0].text = 'Thời gian'
    hdr_cells[1].text = 'Trạng thái hiện tại'
    hdr_cells[2].text = 'Trạng thái tiếp theo'
    for i, ((current, next_state), time) in enumerate(zip(flrs, times[order:])):
        row_cells = flr_table.rows[i + 1].cells
        time_str = str(time) if is_year_only else pd.to_datetime(time).strftime('%Y-%m-%d')
        row_cells[0].text = time_str
        row_cells[1].text = str(current)
        row_cells[2].text = str(next_state)
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_heading('Bảng Nhóm Quan hệ Mờ Phụ thuộc Thời gian (FLRGs) theo Năm', level=1)
    for t, flrgs in enumerate(yearly_flrgs, order):
        time_str = str(times[t]) if is_year_only else pd.to_datetime(times[t]).strftime('%Y-%m-%d')
        doc.add_heading(f'Năm {time_str}', level=2)
        flrg_table = doc.add_table(rows=1 + len(flrgs), cols=3)
        flrg_table.style = 'Table Grid'
        hdr_cells = flrg_table.rows[0].cells
        hdr_cells[0].text = 'Nhóm'
        hdr_cells[1].text = 'Trạng thái hiện tại'
        hdr_cells[2].text = 'Danh sách trạng thái tiếp theo'
        for i, (current, next_states) in enumerate(flrgs.items(), 1):
            row_cells = flrg_table.rows[i].cells
            row_cells[0].text = f'Nhóm {i}'
            row_cells[1].text = str(current)
            row_cells[2].text = str(next_states)
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

    doc.add_heading('Bảng Kết quả Dự báo', level=1)
    result_table = doc.add_table(rows=1 + len(results), cols=4)
    result_table.style = 'Table Grid'
    hdr_cells = result_table.rows[0].cells
    hdr_cells[0].text = 'Thời gian'
    hdr_cells[1].text = 'Giá trị thực'
    hdr_cells[2].text = 'Giá trị dự báo'
    hdr_cells[3].text = 'Tập mờ'
    for i, row in results.iterrows():
        row_cells = result_table.rows[i + 1].cells
        time_str = row['Time'] if is_year_only else pd.to_datetime(row['Time']).strftime('%Y-%m-%d')
        row_cells[0].text = time_str
        row_cells[1].text = str(row['Actual']) if pd.notna(row['Actual']) else ''
        row_cells[2].text = str(round(row['Forecast'], 2)) if pd.notna(row['Forecast']) else ''
        row_cells[3].text = str(row['Fuzzified']) if pd.notna(row['Fuzzified']) else ''
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    try:
        doc.save(output_file)
        print(f"\nĐã xuất bảng ra file Word: {output_file}")
    except Exception as e:
        print(f"Lỗi khi lưu file Word: {e}")

# Hàm chính
def main():
    file_path = input("Nhập đường dẫn file dữ liệu (Excel hoặc CSV): ")
    df, is_year_only = load_data_from_file(file_path)
    if df is None:
        return

    while True:
        try:
            num_intervals = int(input("Nhập số khoảng phân vùng (số nguyên dương): "))
            if num_intervals <= 0:
                print("Số khoảng phải là số nguyên dương!")
                continue
            break
        except ValueError:
            print("Vui lòng nhập một số nguyên hợp lệ!")

    while True:
        try:
            order = int(input("Nhập bậc quan hệ mờ (số nguyên từ 1 đến 9): "))
            if order < 1 or order > 9:
                print("Bậc quan hệ mờ phải từ 1 đến 9!")
                continue
            break
        except ValueError:
            print("Vui lòng nhập một số nguyên hợp lệ!")

    output_file = input("Nhập tên file Word đầu ra (ví dụ: ket_qua_du_bao.docx): ") or "ket_qua_du_bao.docx"
    if not output_file.endswith('.docx'):
        output_file += '.docx'

    data_values = df['Value'].tolist()
    if len(data_values) <= order:
        print(f"Lỗi: Dữ liệu quá ngắn ({len(data_values)} điểm) để dự báo với bậc {order}!")
        return

    D1, D2, optimal_sigmas, mse_opt = optimize_D1_D2_sigmas(data_values, num_intervals, order)

    print("\n" + "=" * 50)
    print(f"D1 tối ưu: {D1:.2f}")
    print(f"D2 tối ưu: {D2:.2f}")
    print(f"Sigmas tối ưu: {[f'{s:.2f}' for s in optimal_sigmas]}")
    print(f"MSE tối ưu: {mse_opt:.0f}")
    print("=" * 50)

    intervals, midpoints = define_universe_and_intervals(data_values, D1, D2, num_intervals)
    display_intervals_and_distribution(data_values, intervals)

    plot_gaussian_membership(intervals, midpoints, optimal_sigmas)

    fuzzy_sets = define_fuzzy_sets(num_intervals)
    print("\nTập hợp mờ (Fuzzy Sets):")
    for key, value in fuzzy_sets.items():
        print(f"{key}: Khoảng {value + 1} (Midpoint: {midpoints[value]:.2f})")

    fuzzified_data = fuzzify_data_gaussian_weighted(data_values, intervals, fuzzy_sets, optimal_sigmas)
    plot_fuzzified_data(df['Time'], data_values, fuzzified_data, is_year_only)

    print("\nDữ liệu sau khi mờ hóa (Fuzzified Data):")
    for time, actual, fuzzy in zip(df['Time'], data_values, fuzzified_data):
        if is_year_only:
            print(f"{time}: {actual} -> {fuzzy}")
        else:
            print(f"{time.strftime('%Y-%m-%d')}: {actual} -> {fuzzy}")

    flrs = identify_flrs(fuzzified_data, order)
    print("\nMối quan hệ mờ (FLRs):")
    for i, (current, next_state) in enumerate(flrs, 1):
        time = df['Time'][i + order - 1]
        time_str = str(time) if is_year_only else time.strftime('%Y-%m-%d')
        print(f"Quan hệ {i} ({time_str}): {current} -> {next_state}")

    yearly_flrgs = establish_time_dependent_flrgs(fuzzified_data, order)
    print("\nNhóm quan hệ mờ phụ thuộc thời gian (FLRGs) theo năm:")
    for t, flrgs in enumerate(yearly_flrgs, order):
        time = df['Time'][t]
        time_str = str(time) if is_year_only else time.strftime('%Y-%m-%d')
        print(f"\nNăm {time_str}:")
        for i, (current, next_states) in enumerate(flrgs.items(), 1):
            print(f"Nhóm {i}: {current} -> {next_states}")

    forecasts = defuzzify_forecast_time_weighted(fuzzified_data, yearly_flrgs, midpoints, order)

    # Đảm bảo cột Fuzzified có cùng độ dài với các cột khác
    results = pd.DataFrame({
        'Time': df['Time'],
        'Actual': data_values,
        'Forecast': forecasts,
        'Fuzzified': fuzzified_data
    })

    valid_results = results.dropna(subset=['Actual', 'Forecast'])
    mse = np.mean((valid_results['Actual'] - valid_results['Forecast']) ** 2)
    mape = np.mean(np.abs((valid_results['Actual'] - valid_results['Forecast']) / valid_results['Actual'])) * 100

    print("\nKết quả mô hình:")
    if is_year_only:
        results['Time'] = results['Time'].astype(str)
    else:
        results['Time'] = results['Time'].dt.strftime('%Y-%m-%d')
    print(results)
    print(f"\nMSE: {mse:.0f}")
    print(f"MAPE: {mape:.2f}%")

    export_to_word(flrs, yearly_flrgs, results, df['Time'], is_year_only, output_file, order)

    plt.figure(figsize=(12, 6))
    plt.plot(results['Time'], results['Actual'], label='Actual', marker='o')
    plt.plot(results['Time'][order:], results['Forecast'][order:], label='Forecast', marker='x')
    plt.title(f"Dự báo chuỗi thời gian (Bậc {order})")
    plt.xlabel("Thời gian")
    plt.ylabel("Giá trị")
    plt.legend()
    plt.grid(True)
    plt.xticks(rotation=45)
    if is_year_only:
        plt.gca().set_xticks(range(0, len(results['Time']), max(1, len(results['Time']) // 10)))
    else:
        plt.gca().xaxis.set_major_formatter(plt.matplotlib.dates.DateFormatter('%Y-%m-%d'))
    plt.tight_layout()
    plt.show()

# Chạy chương trình
if __name__ == "__main__":
    main()