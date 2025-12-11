import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.cluster import KMeans
from docx import Document
from docx.shared import Pt
import io
from sklearn.metrics import davies_bouldin_score
from docx.shared import Inches

# ============= PSO Class for Optimization =============

class TuyChinhKetHop:
    """PSO để tối ưu tìm số cụm K - Thuật toán tùy chỉnh kết hợp"""
    def __init__(self, ham_muc_tieu, cac_bien, so_hat=20, so_lap=50, trong_so_quán_tinh=0.7, 
                 tham_so_tu_nhan=1.5, tham_so_xap_xi=1.5, hat_giong=42):
        self.ham_muc_tieu = ham_muc_tieu
        self.cac_bien = np.array(cac_bien, dtype=float)
        self.so_hat = so_hat
        self.so_lap = so_lap
        self.trong_so_quán_tinh = trong_so_quán_tinh
        self.tham_so_tu_nhan = tham_so_tu_nhan
        self.tham_so_xap_xi = tham_so_xap_xi
        self.may_tao_so_ngau_nhien = np.random.RandomState(hat_giong)

    def chay(self):
        so_chieu = len(self.cac_bien)
        bien_duoi = self.cac_bien[:, 0]
        bien_tren = self.cac_bien[:, 1]
        vi_tri_hien_tai = self.may_tao_so_ngau_nhien.uniform(bien_duoi, bien_tren, size=(self.so_hat, so_chieu))
        van_toc_hien_tai = self.may_tao_so_ngau_nhien.uniform(-np.abs(bien_tren - bien_duoi), 
                                                              np.abs(bien_tren - bien_duoi), 
                                                              size=(self.so_hat, so_chieu))
        vi_tri_tot_nhat_ca_nhan = vi_tri_hien_tai.copy()
        gia_tri_tot_nhat_ca_nhan = np.array([np.inf] * self.so_hat)
        vi_tri_tot_nhat_toan_the = None
        gia_tri_tot_nhat_toan_the = np.inf

        for lap_hien_tai in range(self.so_lap):
            for i in range(self.so_hat):
                x = vi_tri_hien_tai[i]
                gia_tri_hien_tai = self.ham_muc_tieu(x)
                if gia_tri_hien_tai < gia_tri_tot_nhat_ca_nhan[i]:
                    gia_tri_tot_nhat_ca_nhan[i] = gia_tri_hien_tai
                    vi_tri_tot_nhat_ca_nhan[i] = x.copy()
                if gia_tri_hien_tai < gia_tri_tot_nhat_toan_the:
                    gia_tri_tot_nhat_toan_the = gia_tri_hien_tai
                    vi_tri_tot_nhat_toan_the = x.copy()
            
            so_ngau_nhien_1 = self.may_tao_so_ngau_nhien.rand(self.so_hat, so_chieu)
            so_ngau_nhien_2 = self.may_tao_so_ngau_nhien.rand(self.so_hat, so_chieu)
            van_toc_hien_tai = (self.trong_so_quán_tinh * van_toc_hien_tai + 
                               self.tham_so_tu_nhan * so_ngau_nhien_1 * (vi_tri_tot_nhat_ca_nhan - vi_tri_hien_tai) + 
                               self.tham_so_xap_xi * so_ngau_nhien_2 * (vi_tri_tot_nhat_toan_the - vi_tri_hien_tai))
            
            vi_tri_hien_tai = vi_tri_hien_tai + van_toc_hien_tai
            vi_tri_hien_tai = np.maximum(np.minimum(vi_tri_hien_tai, bien_tren), bien_duoi)
        
        return vi_tri_tot_nhat_toan_the, gia_tri_tot_nhat_toan_the

# ============= Helper Functions from Test.py =============

def tim_k_toi_uu_pso(du_lieu, so_cum_toi_da, so_hat, so_lap):
    """Tìm số khoảng tối ưu bằng PSO sử dụng chỉ số Davies-Bouldin (DBI)"""
    
    # Chuẩn bị dữ liệu 1 lần để không phải reshape lặp lại
    X = np.array(du_lieu).reshape(-1, 1)

    def ham_muc_tieu(x):
        # x là mảng chứa K (số cụm)
        k_float = x[0]
        so_cum = int(np.round(k_float))
        
        # Ràng buộc số cụm trong khoảng cho phép
        so_cum = max(2, min(so_cum, so_cum_toi_da))
        
        try:
            # 1. Chạy K-means
            kmeans = KMeans(n_clusters=so_cum, random_state=42, n_init=3)
            nhan_cum = kmeans.fit_predict(X)
            
            # 2. Tính chỉ số Davies-Bouldin (DBI)
            # DBI càng THẤP thì phân cụm càng TỐT
            # Nó tự động phạt nếu các cụm quá gần nhau (chia quá vụn)
            dbi_score = davies_bouldin_score(X, nhan_cum)
            
            return dbi_score
        except:
            return np.inf
    
    # Khởi tạo PSO
    # Bounds: Từ 2 đến so_cum_toi_da
    cac_bien = [(2.0, float(so_cum_toi_da))]
    
    tui_chim_que_hop = TuyChinhKetHop(
        ham_muc_tieu, 
        cac_bien, 
        so_hat=so_hat, 
        so_lap=so_lap, 
        hat_giong=42
    )
    
    # Chạy PSO
    k_toi_uu_vector, gia_tri_toi_uu = tui_chim_que_hop.chay()
    
    # Lấy kết quả cuối cùng
    so_cum_toi_uu = int(np.round(k_toi_uu_vector[0]))
    so_cum_toi_uu = max(2, min(so_cum_toi_uu, so_cum_toi_da))
    
    return so_cum_toi_uu, gia_tri_toi_uu

def tim_k_toi_uu_kmeans(du_lieu, so_cum_toi_da=10):
    """Tìm số khoảng tối ưu bằng K-means sử dụng Elbow method"""
    do_uot = []
    khoang_k = range(2, so_cum_toi_da + 1)
    
    for k in khoang_k:
        tui_chim_kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
        tui_chim_kmeans.fit(np.array(du_lieu).reshape(-1, 1))
        do_uot.append(tui_chim_kmeans.inertia_)
    
    # Tìm điểm elbow - nơi slope thay đổi nhiều nhất
    che_do_sau = np.diff(do_uot)
    che_do_sau_hai_cap = np.diff(che_do_sau)
    so_cum_toi_uu = np.argmax(che_do_sau_hai_cap) + 2
    
    return max(2, min(so_cum_toi_uu, so_cum_toi_da))

def xac_dinh_tap_hop_van_de_va_khoang(du_lieu, so_khoang):
    """Phân vùng tập hợp luận dựa trên K-means"""
    gia_tri_toi_thieu, gia_tri_toi_da = min(du_lieu), max(du_lieu)
    do_chenh_lech_1 = (gia_tri_toi_da - gia_tri_toi_thieu) * 0.001
    do_chenh_lech_2 = (gia_tri_toi_da - gia_tri_toi_thieu) * 0.001
    van_de = [gia_tri_toi_thieu - do_chenh_lech_1, gia_tri_toi_da + do_chenh_lech_2]
    rong_khoang = (van_de[1] - van_de[0]) / so_khoang
    cac_khoang = [[van_de[0] + i * rong_khoang, van_de[0] + (i + 1) * rong_khoang] for i in range(so_khoang)]
    cac_diem_giua = [(u[0] + u[1]) / 2 for u in cac_khoang]
    return cac_khoang, cac_diem_giua

def xac_dinh_tap_hop_mo(so_khoang):
    """Định nghĩa tập hợp mờ"""
    cac_tap_mo = {f'A_{i+1}': i for i in range(so_khoang)}
    return cac_tap_mo

def mo_hoa_du_lieu_gaussian(du_lieu, cac_khoang, cac_tap_mo, cac_sigma=None):
    """Mờ hóa dữ liệu với hàm Gauss theo công thức J(V) = ∑∑ |(xⱼ - cᵢ)|"""
    cac_gia_tri_mo = []
    cac_diem_giua = [(l + u) / 2 for l, u in cac_khoang]
    
    if cac_sigma is None:
        cac_sigma = [(cac_khoang[i][1] - cac_khoang[i][0]) / 2.355 for i in range(len(cac_khoang))]
    
    for x in du_lieu:
        # Tính membership degree cho từng fuzzy set
        do_thuoc_nhom = [np.exp(-((x - m) ** 2) / (2 * (s ** 2))) for m, s in zip(cac_diem_giua, cac_sigma)]
        tong_do_thuoc = sum(do_thuoc_nhom)
        
        if tong_do_thuoc > 0:
            # Tính weighted average của midpoints
            fx = sum(mu_i * m for mu_i, m in zip(do_thuoc_nhom, cac_diem_giua)) / tong_do_thuoc
        else:
            # Nếu không có membership nào, dùng midpoint gần nhất
            fx = cac_diem_giua[np.argmin([abs(x - m) for m in cac_diem_giua])]
        
        # Tìm fuzzy set gần nhất
        chi_so_gan_nhat = np.argmin([abs(fx - m) for m in cac_diem_giua])
        tap_mo_tot_nhat = [khoa for khoa, gia_tri in cac_tap_mo.items() if gia_tri == chi_so_gan_nhat][0]
        cac_gia_tri_mo.append(tap_mo_tot_nhat)
    
    return cac_gia_tri_mo

def xac_dinh_cac_quan_he_mo(du_lieu_mo, bac):
    """Xác định Fuzzy Logical Relations (FLRs)"""
    cac_quan_he = []
    for i in range(len(du_lieu_mo) - bac):
        trang_thai_hien_tai = tuple(du_lieu_mo[i:i + bac])
        trang_thai_tiep_theo = du_lieu_mo[i + bac]
        cac_quan_he.append((trang_thai_hien_tai, trang_thai_tiep_theo))
    return cac_quan_he

def thiet_lap_nhom_quan_he_mo_phu_thuoc_thoi_gian(du_lieu_mo, bac):
    """Thiết lập nhóm quan hệ mờ phụ thuộc thời gian (Time-Dependent FLRGs)"""
    cac_nhom_quan_he_hang_nam = []
    for t in range(bac, len(du_lieu_mo)):
        cac_flrgs = {}
        for i in range(t - bac + 1):
            trang_thai_hien_tai = tuple(du_lieu_mo[i:i + bac])
            trang_thai_tiep_theo = du_lieu_mo[i + bac]
            if trang_thai_hien_tai not in cac_flrgs:
                cac_flrgs[trang_thai_hien_tai] = []
            cac_flrgs[trang_thai_hien_tai].append(trang_thai_tiep_theo)
        cac_nhom_quan_he_hang_nam.append(cac_flrgs)
    return cac_nhom_quan_he_hang_nam

def giai_mo_va_du_bao_co_trong_so_thoi_gian(du_lieu_mo, cac_nhom_quan_he_hang_nam, cac_diem_giua, bac):
    """Giải mờ và dự báo sử dụng trọng số thời gian"""
    cac_du_bao = [None] * bac
    for t in range(bac, len(du_lieu_mo)):
        trang_thai_hien_tai = tuple(du_lieu_mo[t - bac:t])
        cac_flrgs = cac_nhom_quan_he_hang_nam[t - bac]
        
        if trang_thai_hien_tai in cac_flrgs:
            cac_trang_thai_tiep_theo = cac_flrgs[trang_thai_hien_tai]
            cac_trong_so = list(range(1, len(cac_trang_thai_tiep_theo) + 1))
            cac_chi_so = [int(trang_thai.split('_')[1]) - 1 for trang_thai in cac_trang_thai_tiep_theo]
            cac_gia_tri_giua = [cac_diem_giua[chi_so] for chi_so in cac_chi_so]
            
            tu_so = sum(w * m for w, m in zip(cac_trong_so, cac_gia_tri_giua))
            mau_so = sum(cac_trong_so)
            du_bao = tu_so / mau_so if mau_so > 0 else cac_gia_tri_giua[-1]
            cac_du_bao.append(du_bao)
        else:
            # Nếu không tìm thấy quy tắc, dùng midpoint của trạng thái trước đó
            chi_so = int(du_lieu_mo[t - 1].split('_')[1]) - 1
            cac_du_bao.append(cac_diem_giua[chi_so])
    
    return cac_du_bao

def tinh_cac_chi_so_danh_gia(gia_tri_thuc, gia_tri_du_bao):
    """Tính toán MSE và MAPE"""
    cac_chi_so_hop_le = [i for i, (a, p) in enumerate(zip(gia_tri_thuc, gia_tri_du_bao)) 
                     if a is not None and p is not None]
    if len(cac_chi_so_hop_le) == 0:
        return float('inf'), float('inf')
    
    gia_tri_thuc_hop_le = np.array([gia_tri_thuc[i] for i in cac_chi_so_hop_le])
    gia_tri_du_bao_hop_le = np.array([gia_tri_du_bao[i] for i in cac_chi_so_hop_le])
    
    sai_so_binh_phuong_trung_binh = np.mean((gia_tri_thuc_hop_le - gia_tri_du_bao_hop_le) ** 2)
    sai_so_phan_tram_trung_binh = np.mean(np.abs((gia_tri_thuc_hop_le - gia_tri_du_bao_hop_le) / 
                                                 (gia_tri_thuc_hop_le + 1e-10))) * 100
    return sai_so_binh_phuong_trung_binh, sai_so_phan_tram_trung_binh

def them_bieu_do_vao_word(tai_lieu, figure):
    """Chuyển đổi matplotlib figure thành ảnh và chèn vào Word"""
    if figure is not None:
        memfile = io.BytesIO()
        try:
            figure.savefig(memfile, format='png', bbox_inches='tight', dpi=100)
            memfile.seek(0)
            tai_lieu.add_picture(memfile, width=Inches(6)) # Chiều rộng ảnh 6 inches
            tai_lieu.add_paragraph("") # Thêm dòng trống
        except Exception as e:
            tai_lieu.add_paragraph(f"[Lỗi không thể chèn biểu đồ: {e}]")

def xuat_ket_qua_ra_file_word(ket_qua_storage, df_goc, config_info, step1_info, global_figs=None, is_year_only=False):
    """
    Xuất báo cáo Word ĐẦY ĐỦ các bước và biểu đồ.
    """
    tai_lieu = Document()
    style = tai_lieu.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # === TRANG BÌA & TỔNG QUAN ===
    tai_lieu.add_heading('BÁO CÁO DỰ BÁO CHUỖI THỜI GIAN MỜ', 0)
    tai_lieu.add_paragraph(f"Ngày xuất: {pd.Timestamp.now().strftime('%d/%m/%Y %H:%M')}")
    
    # --- 1. Cấu hình mô hình ---
    tai_lieu.add_heading('1. Cấu hình mô hình', level=1)
    p = tai_lieu.add_paragraph()
    p.add_run(f"- K tối đa: {config_info['k_max']}\n")
    p.add_run(f"- PSO (Số hạt): {config_info['pso_pop']}\n")
    p.add_run(f"- PSO (Số vòng lặp): {config_info['pso_iter']}\n")

    # --- 2. Dữ liệu gốc ---
    tai_lieu.add_heading('2. Dữ liệu gốc', level=1)
    # Biểu đồ gốc
    if global_figs and 'fig_goc' in global_figs:
        tai_lieu.add_paragraph("Biểu đồ chuỗi thời gian gốc:")
        them_bieu_do_vao_word(tai_lieu, global_figs['fig_goc'])
    
    # Bảng dữ liệu gốc (15 dòng đầu)
    tai_lieu.add_paragraph("Bảng dữ liệu gốc (15 dòng đầu):")
    t = tai_lieu.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = 'Thời gian'
    t.rows[0].cells[1].text = 'Giá trị'
    for _, row in df_goc.head(15).iterrows():
        r = t.add_row().cells
        r[0].text = str(row['Time']) if is_year_only else pd.to_datetime(row['Time']).strftime('%Y-%m-%d')
        r[1].text = f"{row['Value']:.2f}"

    # --- 3. Bước 1: Xác định tập nền ---
    tai_lieu.add_heading('3. Bước 1: Xác định tập nền (Universe of Discourse)', level=1)
    p = tai_lieu.add_paragraph()
    p.add_run(f"- Min dữ liệu: {step1_info['min_val']:.2f}\n")
    p.add_run(f"- Max dữ liệu: {step1_info['max_val']:.2f}\n")
    p.add_run(f"- Tập nền U = [{step1_info['u_min']:.2f}, {step1_info['u_max']:.2f}]\n").bold = True
    p.add_run(f"- Số cụm tối ưu (K) tìm được: {step1_info['so_cum_toi_uu']}")
    
    if global_figs and 'fig_kmeans' in global_figs:
        tai_lieu.add_heading('Phân cụm K-means tối ưu:', level=2)
        them_bieu_do_vao_word(tai_lieu, global_figs['fig_kmeans'])

    # === CHI TIẾT TỪNG BẬC ===
    cac_bac = list(ket_qua_storage.keys())
    for bac in cac_bac:
        du_lieu = ket_qua_storage[bac]
        tai_lieu.add_page_break()
        
        tai_lieu.add_heading(f'PHẦN KẾT QUẢ CHI TIẾT: BẬC {bac}', level=1)

        # --- Bước 3: Khoảng phân vùng ---
        tai_lieu.add_heading(f'Bước 3: Định nghĩa khoảng phân vùng (Bậc {bac})', level=2)
        if 'fig_phan_bo' in du_lieu:
            them_bieu_do_vao_word(tai_lieu, du_lieu['fig_phan_bo'])
        
        if du_lieu['bang_khoang'] is not None:
            tai_lieu.add_paragraph("Bảng thông số các khoảng:")
            bk = du_lieu['bang_khoang']
            t = tai_lieu.add_table(rows=1, cols=len(bk.columns))
            t.style = 'Table Grid'
            for i, col in enumerate(bk.columns): t.rows[0].cells[i].text = str(col)
            for _, row in bk.iterrows():
                r = t.add_row().cells
                for i, val in enumerate(row): r[i].text = str(val)

        # --- Bước 4: Mờ hóa ---
        tai_lieu.add_heading(f'Bước 4: Mờ hóa dữ liệu (Bậc {bac})', level=2)
        if 'fig_membership' in du_lieu:
            them_bieu_do_vao_word(tai_lieu, du_lieu['fig_membership'])
            
        if du_lieu['bang_mo'] is not None:
            tai_lieu.add_paragraph("Bảng dữ liệu đã mờ hóa (20 dòng đầu):")
            bm = du_lieu['bang_mo'].head(20)
            t = tai_lieu.add_table(rows=1, cols=len(bm.columns))
            t.style = 'Table Grid'
            for i, col in enumerate(bm.columns): t.rows[0].cells[i].text = str(col)
            for _, row in bm.iterrows():
                r = t.add_row().cells
                r[0].text = str(row['Thời gian']) if is_year_only else pd.to_datetime(row['Thời gian']).strftime('%Y-%m-%d')
                r[1].text = str(row['Giá trị'])
                r[2].text = str(row['Tập mờ'])

        # --- Bước 5: Quan hệ mờ (FLRs) ---
        tai_lieu.add_heading(f'Bước 5: Quan hệ mờ (FLRs) (Bậc {bac})', level=2)
        cac_quan_he = du_lieu['cac_quan_he']
        tai_lieu.add_paragraph(f"Tổng số quan hệ: {len(cac_quan_he)}")
        
        t = tai_lieu.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        t.rows[0].cells[0].text = 'Thời gian'
        t.rows[0].cells[1].text = 'Hiện tại'
        t.rows[0].cells[2].text = 'Tiếp theo'
        
        # In tối đa 30 dòng
        for i in range(min(30, len(cac_quan_he))):
            curr, next_val = cac_quan_he[i]
            # Tính lại thời gian tương ứng (dựa vào index + bac)
            # Lưu ý: cần truyền list thoi_gian vào hoặc lấy từ bang_mo
            # Ở đây ta lấy từ bảng mờ cho tiện (dòng i + bac)
            if i + bac < len(du_lieu['bang_mo']):
                time_val = du_lieu['bang_mo'].iloc[i+bac]['Thời gian']
                time_str = str(time_val) if is_year_only else pd.to_datetime(time_val).strftime('%Y-%m-%d')
            else:
                time_str = "N/A"
            
            r = t.add_row().cells
            r[0].text = time_str
            r[1].text = str(curr)
            r[2].text = str(next_val)

        # --- Bước 6: Nhóm quan hệ mờ (FLRGs) ---
        tai_lieu.add_heading(f'Bước 6: Nhóm quan hệ mờ (Bậc {bac})', level=2)
        cac_nhom = du_lieu['cac_nhom_quan_he_hang_nam']
        tai_lieu.add_paragraph("(Hiển thị mẫu 5 thời điểm đầu tiên)")
        
        for idx in range(min(5, len(cac_nhom))):
            # Lấy thời gian
            if idx + bac < len(du_lieu['bang_mo']):
                time_val = du_lieu['bang_mo'].iloc[idx+bac]['Thời gian']
                time_str = str(time_val) if is_year_only else pd.to_datetime(time_val).strftime('%Y-%m-%d')
            else: time_str = "N/A"
            
            p = tai_lieu.add_paragraph()
            p.add_run(f"Thời gian: {time_str}").bold = True
            
            groups = cac_nhom[idx]
            if not groups:
                tai_lieu.add_paragraph("  - Không có nhóm quan hệ.")
            else:
                for gr_idx, (k, v) in enumerate(groups.items(), 1):
                    tai_lieu.add_paragraph(f"  - Nhóm {gr_idx}: {k} -> {v}")

        # --- Giải mờ & Dự báo ---
        tai_lieu.add_heading(f'Bước 7 & 8: Giải mờ và Kết quả (Bậc {bac})', level=2)
        
        # Biểu đồ giải mờ chi tiết
        if 'fig_giai_mo' in du_lieu:
            tai_lieu.add_paragraph("Biểu đồ chi tiết quá trình giải mờ:")
            them_bieu_do_vao_word(tai_lieu, du_lieu['fig_giai_mo'])
            
        tai_lieu.add_paragraph("Bảng kết quả dự báo (20 dòng đầu):")
        bkq = du_lieu['bang_ket_qua'].head(20)
        t = tai_lieu.add_table(rows=1, cols=4)
        t.style = 'Table Grid'
        h = t.rows[0].cells
        h[0].text='Thời gian'; h[1].text='Thực tế'; h[2].text='Dự báo'; h[3].text='Tập mờ'
        for _, row in bkq.iterrows():
            r = t.add_row().cells
            r[0].text = str(row['Thời gian']) if is_year_only else pd.to_datetime(row['Thời gian']).strftime('%Y-%m-%d')
            r[1].text = f"{row['Giá trị thực']:.2f}" if pd.notna(row['Giá trị thực']) else ""
            r[2].text = f"{row['Dự báo']:.2f}" if pd.notna(row['Dự báo']) else ""
            r[3].text = str(row['Tập mờ']) if pd.notna(row['Tập mờ']) else ""

        # --- Bước 9: So sánh ---
        tai_lieu.add_heading(f'Bước 9: So sánh hiệu suất (Bậc {bac})', level=2)
        
        p = tai_lieu.add_paragraph()
        p.add_run(f"MSE: {du_lieu['metrics']['mse']:.2f} | MAPE: {du_lieu['metrics']['mape']:.2f}%").bold = True
        
        # Bảng so sánh
        if du_lieu['bang_so_sanh'] is not None:
            bss = du_lieu['bang_so_sanh']
            t = tai_lieu.add_table(rows=1, cols=len(bss.columns))
            t.style = 'Table Grid'
            for i, col in enumerate(bss.columns): t.rows[0].cells[i].text = str(col)
            for _, row in bss.iterrows():
                r = t.add_row().cells
                r[0].text = str(row['Mô hình'])
                r[1].text = f"{row['MSE']:.4f}"
                r[2].text = f"{row['MAPE (%)']:.2f}"
        
        # 2 Biểu đồ Bước 9
        if 'fig_mape' in du_lieu:
            tai_lieu.add_paragraph("Biểu đồ so sánh sai số MAPE:")
            them_bieu_do_vao_word(tai_lieu, du_lieu['fig_mape'])
            
        if 'fig_trend' in du_lieu:
            tai_lieu.add_paragraph("Biểu đồ so sánh xu hướng dự báo:")
            them_bieu_do_vao_word(tai_lieu, du_lieu['fig_trend'])

    buffer = io.BytesIO()
    tai_lieu.save(buffer)
    buffer.seek(0)
    return buffer



# ============= Streamlit app =============

st.set_page_config(page_title='Dự báo Chuỗi thời gian mờ', layout='wide')
st.title('📊 Dự báo Chuỗi thời gian mờ - K-Means + Giải mờ Gauss')

# Initialize session state
if 'data_loaded' not in st.session_state:
    st.session_state['data_loaded'] = False
    st.session_state['df'] = None
    st.session_state['is_year_only'] = False

# ============= Sidebar: File Upload =============
st.sidebar.header('📁 Tải dữ liệu')
uploaded_file = st.sidebar.file_uploader('Chọn file CSV hoặc Excel', type=['csv', 'xlsx', 'xls'])

# Thêm lựa chọn tổng hợp dữ liệu ngay dưới nút upload
tuy_chon_thoi_gian = st.sidebar.selectbox(
    "Chọn mức độ tổng hợp dữ liệu:",
    ["Ngày", "Tháng", "Năm"],
    index=1,
    help="Chọn 'Tháng' hoặc 'Năm' để tính trung bình cộng giá trị theo thời gian tương ứng."
)

if uploaded_file is not None:
    try:
        filename = uploaded_file.name
        if filename.endswith(('.xlsx', '.xls')):
            df_uploaded = pd.read_excel(uploaded_file)
        else:
            df_uploaded = pd.read_csv(uploaded_file)
        
        # 1. Chuẩn hóa tên cột
        if len(df_uploaded.columns) >= 2:
            df_uploaded = df_uploaded.iloc[:, :2]
            df_uploaded.columns = ['Time', 'Value']
            
            # 2. Xử lý thời gian sang Datetime object để tính toán
            # Cố gắng chuyển đổi cột Time sang datetime
            df_uploaded['Time'] = pd.to_datetime(df_uploaded['Time'], errors='coerce')
            
            # Loại bỏ các dòng không convert được thời gian (NaT)
            so_dong_truoc = len(df_uploaded)
            df_uploaded = df_uploaded.dropna(subset=['Time'])
            
            # 3. Thực hiện tổng hợp dữ liệu (Resample) theo lựa chọn
            if "Tháng" in tuy_chon_thoi_gian:
                # Tính trung bình theo tháng (M = Month end frequency)
                df_uploaded = df_uploaded.set_index('Time').resample('M').mean().reset_index()
                st.sidebar.info("Đã tính trung bình theo Tháng.")
                
            elif "Năm" in tuy_chon_thoi_gian:
                # Tính trung bình theo năm (Y = Year end frequency)
                df_uploaded = df_uploaded.set_index('Time').resample('Y').mean().reset_index()
                st.sidebar.info("Đã tính trung bình theo Năm.")
            
            # 4. Kiểm tra lại logic hiển thị (is_year_only)
            # Biến này dùng để định dạng hiển thị trong biểu đồ/bảng sau này
            is_year_only = False
            if "Năm" in tuy_chon_thoi_gian:
                is_year_only = True
            else:
                # Nếu dữ liệu gốc thực sự chỉ là năm (ví dụ input là integer năm)
                # Logic cũ để kiểm tra fallback
                for t in df_uploaded['Time']:
                    try:
                        # Kiểm tra xem có phải là datetime không, nếu là datetime thì format
                        if isinstance(t, pd.Timestamp):
                            if t.month == 1 and t.day == 1 and "Năm" in tuy_chon_thoi_gian:
                                is_year_only = True
                        else:
                            # Nếu raw data là số
                            int(t)
                            is_year_only = True
                    except (ValueError, TypeError):
                        is_year_only = False
                    break

            # Làm sạch lần cuối
            so_dong_sau = len(df_uploaded)
            if so_dong_truoc > so_dong_sau:
                st.sidebar.warning(f'⚠️ Đã tự động xóa {so_dong_truoc - so_dong_sau} dòng dữ liệu lỗi thời gian.')
            
            df_uploaded = df_uploaded.reset_index(drop=True)
            
            # Lưu vào session_state
            st.session_state['df'] = df_uploaded
            st.session_state['is_year_only'] = is_year_only
            st.session_state['data_loaded'] = True
            st.sidebar.success(f'✓ Tải file thành công: {filename}')
            
            # Hiển thị preview nhỏ kích thước dữ liệu sau khi gộp
            st.sidebar.markdown(f"**Dữ liệu sau xử lý:** {len(df_uploaded)} dòng")
            
        else:
            st.sidebar.error("File phải có ít nhất 2 cột")
    except Exception as e:
        st.sidebar.error(f'❌ Lỗi: {e}')

# ============= Main Content =============
if st.session_state['data_loaded'] and st.session_state['df'] is not None:
    df = st.session_state['df']
    is_year_only = st.session_state['is_year_only']
    data_values = df['Value'].tolist()
    
    # Display raw data
    st.subheader('📊 Dữ liệu gốc')
    col1, col2 = st.columns([3, 1])
    with col1:
        st.dataframe(df.head(10), use_container_width=True)
    with col2:
        st.metric('Số điểm dữ liệu', len(data_values))
        st.metric('Giá trị min', f'{min(data_values):.2f}')
        st.metric('Giá trị max', f'{max(data_values):.2f}')
    
    # Plot original time series
    st.subheader('📈 Biểu đồ chuỗi thời gian gốc')
    fig_goc, ax = plt.subplots(figsize=(14, 5))
    ax.plot(df['Time'], data_values, marker='o', linewidth=2, color='steelblue', markersize=4)
    ax.set_xlabel('Thời gian', fontsize=11)
    ax.set_ylabel('Giá trị', fontsize=11)
    ax.set_title('Chuỗi thời gian gốc', fontsize=12, fontweight='bold')
    ax.grid(True, alpha=0.3)
    plt.xticks(rotation=45)
    plt.tight_layout()
    st.pyplot(fig_goc)
    st.session_state['fig_goc'] = fig_goc
    
    # ============= Model Configuration =============
    st.subheader('⚙️ Cấu hình mô hình')
    
    # Cấu hình PSO cho tìm kiếm số cụm K tối ưu
    col1, col2, col3 = st.columns(3)
    with col1:
        # Slider để chọn K tối đa (giới hạn trên của không gian tìm kiếm)
        so_cum_toi_da = st.slider('Khoảng tối đa:', 2, 30, 7, 
                          help='Số khoảng tối đa để tìm kiếm - PSO sẽ tìm K trong khoảng [2, K tối đa]')
    
    with col2:
        # Số hạt trong đàn PSO - càng lớn càng khám phá kỹ nhưng chậm hơn
        so_hat_pso = st.slider('PSO - Số hạt:', 5, 30, 15,
                              help='Số lượng giải pháp ứng cử trong mỗi vòng lặp')
    
    with col3:
        # Số vòng lặp của PSO - càng lớn càng hội tụ tốt nhưng chậm hơn
        so_lap_pso = st.slider('PSO - Số vòng lặp:', 10, 50, 30,
                              help='Số thế hệ để PSO chạy - quyết định độ hội tụ')
    
    col4, col5 = st.columns(2)
    with col4:
        cac_bac_chon = st.multiselect('Chọn bậc quan hệ mờ:', [1, 3], default=[1, 3],
                                        help='Chọn 1 và/hoặc 3 để chạy mô hình')
    
    if not cac_bac_chon:
        st.warning('⚠️ Vui lòng chọn ít nhất một bậc quan hệ mờ')
    
    # ============= Run Model =============
    if st.button('🔥 Chạy mô hình', key='run_model_btn'):
        if not cac_bac_chon:
            st.error('❌ Vui lòng chọn bậc quan hệ mờ trước')
        else:

# ===== Bước 1: Xác định tập nền =====
            gia_tri_min_goc = min(data_values)
            gia_tri_max_goc = max(data_values)
            do_chenh_lech = (gia_tri_max_goc - gia_tri_min_goc) * 0.1 
            tap_nen_min = gia_tri_min_goc - do_chenh_lech
            tap_nen_max = gia_tri_max_goc + do_chenh_lech
            st.subheader('Bước 1: Xác định tập nền (Universe of Discourse)')
            col_u1, col_u2 = st.columns(2)
            with col_u1:
                st.info(f"**Dữ liệu gốc:**\n\n- Min: {gia_tri_min_goc:.2f}\n- Max: {gia_tri_max_goc:.2f}")
            with col_u2:
                st.success(f"**Tập nền U = [Dmin, Dmax]:**\n\n- Dmin = {tap_nen_min:.2f}\n- Dmax = {tap_nen_max:.2f}\n")
            # ===== Tìm số cụm K tối ưu bằng PSO =====
            with st.spinner('⏳ Tối ưu hóa bằng PSO để tìm số cụm K tối ưu...'):
                st.info(f'🔍 Chạy PSO: {so_hat_pso} hạt, {so_lap_pso} vòng lặp, K ∈ [2, {so_cum_toi_da}]')
                # Sử dụng PSO để tìm K optimal
                so_cum_toi_uu, gia_tri_j = tim_k_toi_uu_pso(data_values, so_cum_toi_da=so_cum_toi_da, 
                                                        so_hat=so_hat_pso, so_lap=so_lap_pso)
                st.success(f'✓ Số cụm tối ưu (PSO): **{so_cum_toi_uu}** | J(V) = {gia_tri_j:.2f}')
            
# ===== Bước 2: Áp dụng K-means với K tối ưu =====
            st.subheader('📊 Bước 2: Áp dụng K-means với K tối ưu tìm được')
            
            # Tính toán J(V) cho các giá trị K khác nhau (để so sánh)
            do_uot = []
            k_thuc_te_max = min(so_cum_toi_da, len(data_values) - 1)
            khoang_k = range(2, k_thuc_te_max + 1)
            
            for k in khoang_k:
                if k >= len(data_values): break # Chặn lỗi nếu vẫn lọt qua
                # Chạy K-means với k cụm
                tui_chim_kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
                tui_chim_kmeans.fit(np.array(data_values).reshape(-1, 1))
                do_uot.append(tui_chim_kmeans.inertia_)
            
            fig_kmeans, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 4))
            
            # Biểu đồ 1: Đường cong J(V) - Công thức K-means
            ax1.plot(khoang_k, do_uot, 'bo-', linewidth=2, markersize=8, label='J(V) cho mỗi K')
            ax1.axvline(x=so_cum_toi_uu, color='red', linestyle='--', linewidth=2.5, 
                       label=f'K tối ưu (PSO) = {so_cum_toi_uu}')
            ax1.set_xlabel('Số cụm (K)', fontsize=11, fontweight='bold')
            ax1.set_ylabel('Inertia J(V) = ∑∑ |(xⱼ - cᵢ)|²', fontsize=11, fontweight='bold')
            ax1.set_title('Hàm mục tiêu K-means - PSO tìm K tối ưu', fontsize=12, fontweight='bold')
            ax1.grid(True, alpha=0.3)
            ax1.legend(fontsize=10)
            
            # Biểu đồ 2: Phân bố dữ liệu vào các cụm K tối ưu
            tui_chim_kmeans_toi_uu = KMeans(n_clusters=so_cum_toi_uu, random_state=42, n_init=10)
            tui_chim_kmeans_toi_uu.fit(np.array(data_values).reshape(-1, 1))
            cac_cum = tui_chim_kmeans_toi_uu.labels_  # Gán cụm cho mỗi điểm dữ liệu
            
            # Vẽ các điểm dữ liệu theo cụm
            for i in range(so_cum_toi_uu):
                cac_diem_cum = np.array(data_values)[cac_cum == i]  # Lấy các điểm trong cụm i
                ax2.scatter([i]*len(cac_diem_cum), cac_diem_cum, alpha=0.6, s=50, label=f'Cụm {i+1}')
            
            # Vẽ tâm cụm
            ax2.scatter(range(so_cum_toi_uu), tui_chim_kmeans_toi_uu.cluster_centers_, 
                       color='red', marker='*', s=500, edgecolor='black', linewidth=2, label='Tâm cụm (centroid)')
            ax2.set_xlabel('Cụm', fontsize=11, fontweight='bold')
            ax2.set_ylabel('Giá trị dữ liệu', fontsize=11, fontweight='bold')
            ax2.set_title(f'Phân bố {len(data_values)} điểm dữ liệu vào {so_cum_toi_uu} cụm', fontsize=12, fontweight='bold')
            ax2.grid(True, alpha=0.3, axis='y')
            ax2.legend(fontsize=9)
            
            plt.tight_layout()
            st.pyplot(fig_kmeans)
            st.session_state['fig_kmeans'] = fig_kmeans
            
            # ===== Xử lý từng bậc quan hệ mờ =====
            ket_qua_theo_bac = {}  # Lưu kết quả cho mỗi bậc
            
            for bac in sorted(cac_bac_chon):  # Lặp qua mỗi bậc được chọn (1 hoặc 3)
                st.markdown('---')
                st.subheader(f'📍 Bậc quan hệ mờ = {bac}')
                
                # Validate data length
                if len(data_values) <= bac:
                    st.error(f'❌ Dữ liệu quá ngắn ({len(data_values)} điểm) cho bậc {bac}')
                    continue
                
# ===== Bước 3: Định nghĩa khoảng phân vùng =====
                # Sử dụng K tối ưu từ PSO để chia dữ liệu thành K khoảng
                st.subheader('**Bước 3: Định nghĩa khoảng phân vùng (Universe of Discourse)**')
                # Tính toán khoảng phân vùng dựa trên min, max và số cụm
                cac_khoang, cac_diem_giua = xac_dinh_tap_hop_van_de_va_khoang(data_values, so_cum_toi_uu)
                
                # Hiển thị bảng khoảng phân vùng với các biên và trung tâm
                bang_khoang = pd.DataFrame({
                    'Khoảng': [f'K{i+1}' for i in range(len(cac_khoang))],
                    'Biên trái': [f'{u[0]:.2f}' for u in cac_khoang],
                    'Trung tâm (Midpoint)': [f'{m:.2f}' for m in cac_diem_giua],
                    'Biên phải': [f'{u[1]:.2f}' for u in cac_khoang]
                })
                st.dataframe(bang_khoang, use_container_width=True)
                
                # Trực quan hóa phân bố dữ liệu vào các khoảng
                # Đếm số điểm dữ liệu trong mỗi khoảng
                phan_bo = [0] * len(cac_khoang)
                for gia_tri in data_values:
                    for i, (tren_duoi, tren_tren) in enumerate(cac_khoang):
                        # Nếu giá trị nằm trong khoảng [tren_duoi, tren_tren]
                        if tren_duoi <= gia_tri <= tren_tren:
                            phan_bo[i] += 1
                            break
                
                fig_phan_bo, ax = plt.subplots(figsize=(12, 4))
                # Vẽ biểu đồ cột thể hiện số điểm trong mỗi khoảng
                ax.bar(range(1, len(cac_khoang) + 1), phan_bo, color='steelblue', alpha=0.7, edgecolor='black')
                ax.set_xlabel('Khoảng (Interval)', fontsize=11, fontweight='bold')
                ax.set_ylabel('Số lượng giá trị', fontsize=11, fontweight='bold')
                ax.set_title(f'Phân bố dữ liệu trong {so_cum_toi_uu} khoảng (Bậc {bac})', fontsize=12, fontweight='bold')
                ax.grid(True, alpha=0.3, axis='y')
                plt.xticks(range(1, len(cac_khoang) + 1))
                plt.tight_layout()
                st.pyplot(fig_phan_bo)
                
# ===== Bước 4: Mờ hóa dữ liệu =====
                # Dùng hàm Gauss để gán độ thuộc cho mỗi điểm dữ liệu
                st.subheader('**Bước 4: Mờ hóa dữ liệu (Fuzzification)**')
                # Định nghĩa các tập hợp mờ (fuzzy sets)
                cac_tap_mo = xac_dinh_tap_hop_mo(so_cum_toi_uu)
                # Tính độ lệch chuẩn cho hàm Gauss từ chiều rộng khoảng
                cac_sigma = [(cac_khoang[i][1] - cac_khoang[i][0]) / 2.355 for i in range(len(cac_khoang))]
                # Mờ hóa dữ liệu sử dụng hàm Gauss
                du_lieu_mo = mo_hoa_du_lieu_gaussian(data_values, cac_khoang, cac_tap_mo, cac_sigma)
                
                bang_mo = pd.DataFrame({
                    'Thời gian': df['Time'],
                    'Giá trị': [f'{v:.2f}' for v in data_values],
                    'Tập mờ': du_lieu_mo
                })
                st.dataframe(bang_mo.head(15), use_container_width=True)
                
                # Visualize membership functions
                fig_membership, ax = plt.subplots(figsize=(12, 4))
                x_values = np.linspace(min([l for l, _ in cac_khoang]) - 10, max([h for _, h in cac_khoang]) + 10, 500)
                for i, (m, s) in enumerate(zip(cac_diem_giua, cac_sigma)):
                    mu = [np.exp(-((x - m) ** 2) / (2 * (s ** 2))) for x in x_values]
                    ax.plot(x_values, mu, label=f'A_{i+1} (c={m:.1f})', linewidth=2)
                ax.set_title(f'Hàm membership Gauss (Bậc {bac})', fontsize=12, fontweight='bold')
                ax.set_xlabel('Giá trị', fontsize=11)
                ax.set_ylabel('Membership (μ)', fontsize=11)
                ax.legend(ncol=min(4, len(cac_diem_giua)), fontsize='small')
                ax.grid(True, alpha=0.3)
                ax.set_ylim([0, 1.05])
                plt.tight_layout()
                st.pyplot(fig_membership)
                
# ===== Bước 5: Xác định Quan hệ Mờ (FLRs) =====
                # FLR: Fuzzy Logical Relation - quan hệ giữa các trạng thái mờ
                st.subheader('**Bước 5: Quan hệ mờ (Fuzzy Logical Relations - FLRs)**')
                cac_quan_he = xac_dinh_cac_quan_he_mo(du_lieu_mo, bac)
                
                bang_quan_he = pd.DataFrame({
                    'Thời gian': [str(df['Time'][i + bac]) if is_year_only else pd.to_datetime(df['Time'][i + bac]).strftime('%Y-%m-%d') for i in range(len(cac_quan_he))],
                    'Trạng thái hiện tại': [str(trang_thai_hien_tai) for trang_thai_hien_tai, _ in cac_quan_he],
                    'Trạng thái tiếp theo': [trang_thai_tiep_theo for _, trang_thai_tiep_theo in cac_quan_he]
                })
                st.dataframe(bang_quan_he.head(15), use_container_width=True)
                
# ===== Bước 6: Nhóm Quan hệ Mờ phụ thuộc Thời gian =====
                # FLRG: Fuzzy Logical Relation Group - nhóm các FLR theo thời gian
                st.subheader('Bước 6: Nhóm quan hệ mờ ')
                cac_nhom_quan_he_hang_nam = thiet_lap_nhom_quan_he_mo_phu_thuoc_thoi_gian(du_lieu_mo, bac)
                
                # Display a few FLRGs
                for t_idx in range(min(3, len(cac_nhom_quan_he_hang_nam))):
                    cac_flrgs = cac_nhom_quan_he_hang_nam[t_idx]
                    thoi_gian_diem = df['Time'][t_idx + bac]
                    chuoi_thoi_gian = str(thoi_gian_diem) if is_year_only else pd.to_datetime(thoi_gian_diem).strftime('%Y-%m-%d')
                    
                    with st.expander(f'📅 Thời gian {chuoi_thoi_gian}', expanded=False):
                        for chi_so_nhom, (trang_thai_hien_tai, cac_trang_thai_tiep_theo) in enumerate(cac_flrgs.items(), 1):
                            st.write(f'**Nhóm {chi_so_nhom}:** {trang_thai_hien_tai} → {cac_trang_thai_tiep_theo}')

# ===== Bước 7: Giải mờ =====
                st.subheader('Bước 7: Giải mờ ')
                
                # Thực hiện dự báo
                cac_du_bao = giai_mo_va_du_bao_co_trong_so_thoi_gian(du_lieu_mo, cac_nhom_quan_he_hang_nam, cac_diem_giua, bac)
                
                # Tạo dataframe kết quả
                bang_ket_qua = pd.DataFrame({
                    'Thời gian': df['Time'],
                    'Giá trị thực': data_values,
                    'Dự báo': cac_du_bao,
                    'Tập mờ': du_lieu_mo
                })
                
                # --- THÊM MỚI: Biểu đồ giải mờ chi tiết ---
                st.markdown("##### 📈 Biểu đồ chi tiết quá trình giải mờ")
                
                # Lấy một khoảng dữ liệu mẫu để vẽ cho thoáng (ví dụ 50 điểm cuối hoặc toàn bộ nếu ít)
                so_diem_ve = 50
                if len(bang_ket_qua) > so_diem_ve:
                    df_ve = bang_ket_qua.iloc[-so_diem_ve:].copy()
                    start_idx = len(bang_ket_qua) - so_diem_ve
                else:
                    df_ve = bang_ket_qua.copy()
                    start_idx = 0
                
                fig_gm, ax_gm = plt.subplots(figsize=(14, 6))
                
                # Vẽ đường dự báo và thực tế
                ax_gm.plot(df_ve['Thời gian'], df_ve['Giá trị thực'], label='Thực tế', color='lightgray', linestyle='--', alpha=0.7)
                ax_gm.plot(df_ve['Thời gian'], df_ve['Dự báo'], label='Dự báo (Giải mờ)', color='#1f77b4', marker='o', markersize=4)
                
                # Thêm nhãn tập mờ lên trên các điểm dự báo
                # Chỉ hiện nhãn cho các điểm dự báo có giá trị (không phải None)
                for i in range(len(df_ve)):
                    val = df_ve['Dự báo'].iloc[i]
                    tap_mo = df_ve['Tập mờ'].iloc[i]
                    thoi_gian = df_ve['Thời gian'].iloc[i]
                    
                    if pd.notna(val):
                        # Offset y một chút để chữ nằm trên điểm
                        ax_gm.text(thoi_gian, val + (max(data_values) - min(data_values))*0.02, 
                                 f"{tap_mo}", 
                                 fontsize=8, 
                                 ha='center', 
                                 color='darkred',
                                 rotation=0)
                
                ax_gm.set_title(f'Minh họa kết quả giải mờ và nhãn tập mờ tương ứng (Bậc {bac})', fontweight='bold')
                ax_gm.set_ylabel('Giá trị')
                ax_gm.legend()
                ax_gm.grid(True, alpha=0.3)
                plt.xticks(rotation=45)
                st.pyplot(fig_gm)
                st.session_state[f'fig_gm_{bac}'] = fig_gm
                # ------------------------------------------
                
# ===== Bước 8:Dự báo và đánh giá hiệu suất =====
                st.subheader('**Bước 8: Dự báo và đánh giá độ chính xác**')
                
                # Tính toán các chỉ số đánh giá
                sai_so_bptp, sai_so_mape = tinh_cac_chi_so_danh_gia(data_values, cac_du_bao)
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric('MSE (Sai số bình phương)', f'{sai_so_bptp:.2f}')
                with col2:
                    st.metric('MAPE (Sai số %)', f'{sai_so_mape:.2f}%')
                with col3:
                    st.metric('Số khoảng (Cluster)', f'{so_cum_toi_uu}')
                with col4:
                    st.metric('Bậc quan hệ', f'{bac}')
                
                # --- Bảng kết quả ---
                st.markdown("##### 📋 Bảng dữ liệu kết quả dự báo")
                st.dataframe(bang_ket_qua.head(20), use_container_width=True)

                # ===== Trực quan hóa tổng thể: So sánh Giá trị Thực vs Dự báo =====
                fig_compare, ax = plt.subplots(figsize=(14, 5))
                
                # Xử lý các giá trị dự báo None để vẽ liền mạch hơn
                du_bao_hop_le = [f if f is not None else np.nan for f in cac_du_bao]
                
                # Vẽ đường giá trị thực
                ax.plot(df['Time'], data_values, label='Giá trị thực (Actual)', marker='.', 
                       linewidth=1.5, color='black', alpha=0.6)
                
                # Vẽ đường dự báo
                ax.plot(df['Time'], du_bao_hop_le, label='Dự báo (Forecast)', marker='x', 
                       linewidth=2, color='red', markersize=4)
                       
                ax.set_xlabel('Thời gian', fontsize=11, fontweight='bold')
                ax.set_ylabel('Giá trị', fontsize=11, fontweight='bold')
                ax.set_title(f'So sánh tổng thể: Thực tế vs Dự báo (Bậc {bac}, K={so_cum_toi_uu})', 
                            fontsize=12, fontweight='bold')
                ax.legend(loc='best', fontsize=11)
                ax.grid(True, alpha=0.3)
                plt.xticks(rotation=45)
                plt.tight_layout()
                st.pyplot(fig_compare)
                

# ===== Bước 9: So sánh với thuật toán khác =====
                st.subheader('**Bước 9: So sánh hiệu suất với các mô hình khác**')
                
                from sklearn.linear_model import LinearRegression
                
                # 1. Chuẩn bị dữ liệu
                # ### SỬA LỖI TẠI ĐÂY: Thêm dtype=float để đảm bảo mảng chứa được np.nan
                y_true = np.array(data_values, dtype=float)
                n = len(y_true)
                
                # --- Mô hình 1: Naive ---
                # Dự báo = Giá trị thực của thời điểm trước đó
                y_naive = np.roll(y_true, 1)
                y_naive[0] = np.nan # Bây giờ mảng đã là float nên dòng này sẽ không lỗi nữa
                
                # --- Mô hình 2: Simple Moving Average (SMA - MA3) ---
                # Trung bình cộng 3 điểm gần nhất
                y_sma = pd.Series(y_true).rolling(window=3).mean().shift(1).values
                
                # --- Mô hình 3: Linear Regression (AR1) ---
                # Dùng sklearn LinearRegression để dự báo y_t dựa trên y_{t-1}
                # Cần xử lý NaN hoặc cắt bỏ phần tử đầu tiên để fit model
                X_lr = y_true[:-1].reshape(-1, 1) # Feature: Lag 1
                y_lr_target = y_true[1:]          # Target: Current value
                
                if len(X_lr) > 0:
                    reg = LinearRegression().fit(X_lr, y_lr_target)
                    y_lr_pred = reg.predict(y_true.reshape(-1, 1)) # Predict toàn bộ chuỗi
                    # Shift kết quả vì dự báo là cho t+1, nhưng ta đang align với t
                    y_lr_final = np.full(n, np.nan)
                    y_lr_final[1:] = y_lr_pred[:-1]
                else:
                    y_lr_final = np.full(n, np.nan)
                
                # --- Hàm tính sai số ---
                def tinh_sai_so_so_sanh(y_t, y_p):
                    # Chỉ tính trên các điểm mà cả 2 đều không phải NaN
                    mask = ~np.isnan(y_t) & ~np.isnan(y_p)
                    if np.sum(mask) == 0: return np.inf, np.inf
                    mse = np.mean((y_t[mask] - y_p[mask])**2)
                    mape = np.mean(np.abs((y_t[mask] - y_p[mask]) / (y_t[mask] + 1e-10))) * 100
                    return mse, mape

                # Chuẩn bị vector FTS (xử lý None thành NaN để tính toán)
                y_fts_calc = np.array([x if x is not None else np.nan for x in cac_du_bao], dtype=float)
                
                # Tính sai số cho từng mô hình
                mse_fts, mape_fts = tinh_sai_so_so_sanh(y_true, y_fts_calc)
                mse_naive, mape_naive = tinh_sai_so_so_sanh(y_true, y_naive)
                mse_sma, mape_sma = tinh_sai_so_so_sanh(y_true, y_sma)
                mse_lr, mape_lr = tinh_sai_so_so_sanh(y_true, y_lr_final)
                
                # --- Hiển thị bảng so sánh ---
                bang_so_sanh = pd.DataFrame({
                    'Mô hình': ['Fuzzy Time Series (PSO + K-means + Gauss)', 'Naive (1-step ahead)', 'Moving Average (SMA-3)', 'Linear Regression (AR1)'],
                    'MSE': [mse_fts, mse_naive, mse_sma, mse_lr],
                    'MAPE (%)': [mape_fts, mape_naive, mape_sma, mape_lr]
                })
                
                # Highlight mô hình tốt nhất (MAPE thấp nhất)
                st.markdown("##### 📊 Bảng so sánh độ chính xác (MAPE thấp hơn là tốt hơn)")
                st.dataframe(bang_so_sanh.style.highlight_min(axis=0, subset=['MSE', 'MAPE (%)'], color='red'), use_container_width=True)
                
                # --- Biểu đồ so sánh MAPE ---
                col_chart1, col_chart2 = st.columns(2)
                
                with col_chart1:
                    fig_bar, ax_bar = plt.subplots(figsize=(6, 4))
                    colors = ['#ff4b4b', 'gray', 'gray', 'gray'] 
                    ax_bar.bar(bang_so_sanh['Mô hình'], bang_so_sanh['MAPE (%)'], color=colors, alpha=0.8)
                    ax_bar.set_ylabel('MAPE (%)')
                    ax_bar.set_title('So sánh sai số MAPE', fontweight='bold')
                    plt.xticks(rotation=45, ha='right')
                    plt.grid(axis='y', alpha=0.3)
                    st.pyplot(fig_bar)
                    fig_mape = fig_bar

                with col_chart2:
                    # --- Biểu đồ đường so sánh các đường dự báo ---
                    fig_comp, ax_comp = plt.subplots(figsize=(8, 4))
                    # Vẽ dữ liệu thực tế
                    # Chỉ vẽ tối đa 50 điểm cuối để dễ nhìn
                    so_diem_ve_ss = 50
                    start_idx_ss = max(0, len(y_true) - so_diem_ve_ss)
                    
                    ax_comp.plot(df['Time'][start_idx_ss:], y_true[start_idx_ss:], label='Thực tế', color='black', linewidth=2, linestyle='-')
                    # Vẽ FTS
                    ax_comp.plot(df['Time'][start_idx_ss:], y_fts_calc[start_idx_ss:], label='Fuzzy Time Series (PSO + K-means + Gauss)', color='#ff4b4b', linewidth=2)
                    # Vẽ Naive
                    ax_comp.plot(df['Time'][start_idx_ss:], y_naive[start_idx_ss:], label='Naive', color='green', linestyle='--', alpha=0.6)
                    # Vẽ SMA
                    ax_comp.plot(df['Time'][start_idx_ss:], y_sma[start_idx_ss:], label='SMA-3', color='blue', linestyle=':', alpha=0.6)
                    
                    ax_comp.set_title(f'So sánh xu hướng dự báo ({len(y_true[start_idx_ss:])} điểm cuối)', fontweight='bold')
                    ax_comp.legend(fontsize='small')
                    ax_comp.grid(True, alpha=0.3)
                    plt.xticks(rotation=45)
                    st.pyplot(fig_comp)
                    fig_trend = fig_comp

            ket_qua_theo_bac[bac] = {
                    'cac_khoang': cac_khoang,
                    'bang_khoang': bang_khoang,
                    'bang_mo': bang_mo,
                    'cac_quan_he': cac_quan_he,
                    'cac_nhom_quan_he_hang_nam': cac_nhom_quan_he_hang_nam,
                    'bang_ket_qua': bang_ket_qua,
                    'bang_so_sanh': bang_so_sanh,
                    'metrics': {'mse': sai_so_bptp, 'mape': sai_so_mape},
                    'so_cum_toi_uu': so_cum_toi_uu,
                    'fig_phan_bo': fig_phan_bo,      
                    'fig_membership': fig_membership,
                    'fig_compare': fig_compare,    
                    'fig_giai_mo': fig_gm,           
                    'fig_mape': fig_mape,         
                    'fig_trend': fig_trend           
                }
                # ===== Xuất kết quả ra file Word ===== 
            st.session_state['results_storage'] = ket_qua_theo_bac 
            st.success("✅ Đã chạy xong mô hình! Kéo xuống dưới để xuất file.")

    # ============= PHẦN XUẤT FILE (CẬP NHẬT) =============
    if 'results_storage' in st.session_state and st.session_state['results_storage']:
        st.markdown('---')
        st.subheader('📥 Xuất báo cáo tổng hợp')
        
        toan_bo_ket_qua = st.session_state['results_storage']
        ds_bac = "_".join(map(str, toan_bo_ket_qua.keys()))
        ten_file_xuat = f"Bao_cao_Chi_tiet_{ds_bac}.docx"

        first_key = list(toan_bo_ket_qua.keys())[0]
        so_cum = toan_bo_ket_qua[first_key]['so_cum_toi_uu']
        
        vals = st.session_state['df']['Value'].tolist()
        min_v, max_v = min(vals), max(vals)
        diff = (max_v - min_v) * 0.1
        
        step1_info = {
            'min_val': min_v,
            'max_val': max_v,
            'u_min': min_v - diff,
            'u_max': max_v + diff,
            'so_cum_toi_uu': so_cum
        }
        
        config_info = {
            'k_max': so_cum_toi_da, 
            'pso_pop': so_hat_pso,  
            'pso_iter': so_lap_pso 
        }
        
        global_figs = {
            'fig_goc': st.session_state.get('fig_goc'),
            'fig_kmeans': st.session_state.get('fig_kmeans')
        }
        
        file_docx = xuat_ket_qua_ra_file_word(
            ket_qua_storage=toan_bo_ket_qua,
            df_goc=st.session_state['df'],    
            config_info=config_info,          
            step1_info=step1_info,           
            global_figs=global_figs,
            is_year_only=st.session_state['is_year_only']
        )
        
        st.download_button(
            label="💾 Tải xuống Báo cáo Full (Word)",
            data=file_docx,
            file_name=ten_file_xuat,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

else:
    st.info('👈 Vui lòng tải file dữ liệu từ thanh bên trái để bắt đầu')
    st.markdown("""
    ### 📋 Hướng dẫn sử dụng:
    
    1. **Tải dữ liệu**: Chọn file CSV hoặc Excel từ thanh bên trái
       - File phải có 2 cột: Thời gian (Time) và Giá trị (Value)
       - Định dạng thời gian: Năm (YYYY) hoặc Ngày (YYYY-MM-DD)
    
    2. **Cấu hình mô hình**:
       - Chọn K tối đa (mặc định: 14)
       - **Phương pháp tối ưu**:
       - Chọn PSO - Số hạt (mặc định: 15) - càng lớn càng khám phá kỹ nhưng chất lượng tìm kiếm cơ bản
       - Chọn PSO - Số vòng lặp (mặc định: 30) - càng lớn càng hội tụ tốt nhưng chậm hơn
    
    
    3. **Chạy mô hình**:
       - Nhấn nút "🔥 Chạy mô hình"
       - Hệ thống sẽ tự động tìm số cụm tối ưu
       - Hiển thị 5 bước xử lý: Định nghĩa khoảng, Mờ hóa, FLRs, FLRGs, Dự báo
    
    4. **Xuất kết quả**:
       - Xuất kết quả ra file Word (*.docx)
       - Bao gồm các bảng FLRs, FLRGs, và kết quả dự báo
    
    """)
